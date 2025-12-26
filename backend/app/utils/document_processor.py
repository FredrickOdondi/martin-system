"""
Document Processing Utilities

This module provides utilities for processing various document formats,
extracting text, and preparing documents for embedding and indexing.
"""

from typing import List, Dict, Any, Optional, Tuple
import os
import re
from pathlib import Path
import logging
from datetime import datetime

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes documents for ingestion into the knowledge base.
    
    Supports:
    - PDF files
    - Word documents (.docx)
    - Text files (.txt, .md)
    - Excel files (.xlsx, .csv)
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'text',
        '.md': 'markdown',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv'
    }
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_chunks_per_doc: int = 100
    ):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Target size for text chunks (in tokens)
            chunk_overlap: Overlap between chunks (in tokens)
            max_chunks_per_doc: Maximum chunks per document
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chunks_per_doc = max_chunks_per_doc
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if supported, False otherwise
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from document based on file type.
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text content
        """
        ext = Path(file_path).suffix.lower()
        
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file type: {ext}")
        
        file_type = self.SUPPORTED_EXTENSIONS[ext]
        
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            elif file_type in ['text', 'markdown']:
                return self._extract_from_text(file_path)
            elif file_type in ['excel', 'csv']:
                return self._extract_from_excel(file_path)
            else:
                raise ValueError(f"Handler not implemented for: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        return '\n\n'.join(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        doc = DocxDocument(file_path)
        
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)
        
        return '\n\n'.join(text)
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel or CSV file."""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Convert dataframe to text representation
        text = df.to_string(index=False)
        return text
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove multiple spaces
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\-\:\;\(\)\[\]\{\}\"\'\/]', '', text)
        
        # Strip whitespace
        text = text.strip()
        
        return text
    
    def chunk_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks with overlap.
        
        Args:
            text: Text to chunk
            metadata: Optional metadata to attach to each chunk
            
        Returns:
            List of chunks with metadata
        """
        # Simple word-based chunking (can be improved with tiktoken)
        words = text.split()
        chunks = []
        
        # Calculate chunk sizes in words (approximate)
        words_per_chunk = self.chunk_size // 4  # Rough estimate: 1 token ≈ 0.75 words
        overlap_words = self.chunk_overlap // 4
        
        for i in range(0, len(words), words_per_chunk - overlap_words):
            chunk_words = words[i:i + words_per_chunk]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text.strip()) > 50:  # Minimum chunk size
                chunk = {
                    'text': chunk_text,
                    'metadata': {
                        **(metadata or {}),
                        'chunk_index': len(chunks),
                        'total_chunks': 0  # Will be updated later
                    }
                }
                chunks.append(chunk)
            
            # Limit number of chunks
            if len(chunks) >= self.max_chunks_per_doc:
                logger.warning(f"Reached max chunks limit: {self.max_chunks_per_doc}")
                break
        
        # Update total_chunks in metadata
        for chunk in chunks:
            chunk['metadata']['total_chunks'] = len(chunks)
        
        return chunks
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dict with file metadata
        """
        path = Path(file_path)
        stat = path.stat()
        
        metadata = {
            'filename': path.name,
            'file_type': path.suffix.lower(),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }
        
        return metadata
    
    def process_document(
        self,
        file_path: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Complete document processing pipeline.
        
        Args:
            file_path: Path to document
            additional_metadata: Additional metadata to attach
            
        Returns:
            Dict with processed document data
        """
        logger.info(f"Processing document: {file_path}")
        
        try:
            # Extract text
            raw_text = self.extract_text(file_path)
            
            # Clean text
            cleaned_text = self.clean_text(raw_text)
            
            # Extract file metadata
            file_metadata = self.extract_metadata(file_path)
            
            # Merge metadata
            metadata = {
                **file_metadata,
                **(additional_metadata or {})
            }
            
            # Chunk text
            chunks = self.chunk_text(cleaned_text, metadata)
            
            result = {
                'file_path': file_path,
                'raw_text': raw_text,
                'cleaned_text': cleaned_text,
                'chunks': chunks,
                'metadata': metadata,
                'chunk_count': len(chunks),
                'status': 'success'
            }
            
            logger.info(f"Processed {file_path}: {len(chunks)} chunks created")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                'file_path': file_path,
                'status': 'failed',
                'error': str(e)
            }
    
    def batch_process(
        self,
        file_paths: List[str],
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of file paths
            additional_metadata: Metadata to attach to all documents
            
        Returns:
            List of processed document results
        """
        results = []
        
        for file_path in file_paths:
            result = self.process_document(file_path, additional_metadata)
            results.append(result)
        
        successful = sum(1 for r in results if r['status'] == 'success')
        logger.info(f"Batch processed {len(file_paths)} documents: {successful} successful")
        
        return results


def get_document_processor() -> DocumentProcessor:
    """
    Get DocumentProcessor instance with default settings.
    
    Returns:
        DocumentProcessor instance
    """
    chunk_size = int(os.getenv("CHUNK_SIZE", "1000"))
    chunk_overlap = int(os.getenv("CHUNK_OVERLAP", "200"))
    max_chunks = int(os.getenv("MAX_CHUNKS_PER_DOC", "100"))
    
    return DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        max_chunks_per_doc=max_chunks
    )
